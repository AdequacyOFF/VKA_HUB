"""Report generation using python-docx"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import get_settings

settings = get_settings()


def generate_docx_report(competition_data: Dict[str, Any]) -> str:
    """
    Generate a .docx report for a competition.

    Args:
        competition_data: Dictionary containing:
            - competition_name: str
            - start_date: date
            - end_date: date
            - teams: List[Dict] with team and member data

    Returns:
        File path to generated report
    """
    # Create document
    doc = Document()

    # Add title
    title = doc.add_heading(competition_data.get("competition_name", "Competition Report"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add competition info
    doc.add_paragraph(f"Start Date: {competition_data.get('start_date', 'N/A')}")
    doc.add_paragraph(f"End Date: {competition_data.get('end_date', 'N/A')}")
    doc.add_paragraph(f"Report Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_paragraph()  # Blank line

    # Add teams section
    doc.add_heading("Participating Teams", level=2)

    teams = competition_data.get("teams", [])

    if teams:
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Team Name'
        hdr_cells[1].text = 'Captain'
        hdr_cells[2].text = 'Members'

        # Data rows
        for team in teams:
            row_cells = table.add_row().cells
            row_cells[0].text = team.get("name", "")
            row_cells[1].text = team.get("captain", "")
            row_cells[2].text = ", ".join(team.get("members", []))
    else:
        doc.add_paragraph("No teams registered for this competition.")

    # Save document
    reports_dir = Path(settings.UPLOAD_DIR) / "reports_generated"
    reports_dir.mkdir(parents=True, exist_ok=True)

    filename = f"competition_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = reports_dir / filename

    doc.save(str(file_path))

    return f"/static/uploads/reports_generated/{filename}"
